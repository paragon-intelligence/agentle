import hashlib
from pathlib import Path
from typing import Literal, override

from rsb.functions.ext_to_mime import ext_to_mime
from rsb.models.field import Field

from agentle.agents.agent import Agent
from agentle.generations.models.message_parts.file import FilePart
from agentle.generations.models.structured_outputs_store.visual_media_description import (
    VisualMediaDescription,
)
from agentle.generations.providers.base.generation_provider import GenerationProvider
from agentle.generations.providers.google.google_genai_generation_provider import (
    GoogleGenaiGenerationProvider,
)
from agentle.parsing.document_parser import DocumentParser
from agentle.parsing.factories.visual_description_agent_factory import (
    visual_description_agent_factory,
)
from agentle.parsing.image import Image
from agentle.parsing.parsed_document import ParsedDocument
from agentle.parsing.parses import parses
from agentle.parsing.section_content import SectionContent


@parses("doc", "docx")
class DocxFileParser(DocumentParser):
    strategy: Literal["high", "low"] = Field(default="high")

    visual_description_agent: Agent[VisualMediaDescription] = Field(
        default_factory=visual_description_agent_factory,
    )
    """
    The agent to use for generating the visual description of the document.
    Useful when you want to customize the prompt for the visual description.
    """

    multi_modal_provider: GenerationProvider = Field(
        default_factory=GoogleGenaiGenerationProvider,
    )
    """
    The multi-modal provider to use for generating the visual description of the document.
    Useful when you want us to customize the prompt for the visual description.
    """

    @override
    async def parse_async(
        self,
        document_path: str,
    ) -> ParsedDocument:
        from docx import Document

        document = Document(document_path)
        image_cache: dict[str, tuple[str, str]] = {}  # (md, ocr_text)
        extension = Path(document_path).suffix

        paragraph_texts = [p.text for p in document.paragraphs if p.text.strip()]
        doc_text = "\n".join(paragraph_texts)

        doc_images: list[tuple[str, bytes]] = []
        for rel in document.part._rels.values():  # type: ignore[reportPrivateUsage]
            if "image" in rel.reltype:
                image_part = rel.target_part
                image_name = image_part.partname.split("/")[-1]
                image_bytes = image_part.blob
                doc_images.append((image_name, image_bytes))

        final_images: list[Image] = []
        image_descriptions: list[str] = []
        if self.visual_description_agent and self.strategy == "high":
            for idx, (image_name, image_bytes) in enumerate(doc_images, start=1):
                image_hash = hashlib.sha256(image_bytes).hexdigest()

                if image_hash in image_cache:
                    cached_md, cached_ocr = image_cache[image_hash]
                    image_md = cached_md
                    ocr_text = cached_ocr
                else:
                    agent_input = FilePart(
                        mime_type=ext_to_mime(extension),
                        data=image_bytes,
                    )
                    agent_response = await self.visual_description_agent.run_async(
                        agent_input
                    )
                    image_md = agent_response.parsed.md
                    ocr_text = agent_response.parsed.ocr_text
                    image_cache[image_hash] = (image_md, ocr_text or "")

                image_descriptions.append(f"Docx Image {idx}: {image_md}")
                final_images.append(
                    Image(
                        name=image_name,
                        contents=image_bytes,
                        ocr_text=ocr_text,
                    )
                )

            if image_descriptions:
                doc_text += "\n\n" + "\n".join(image_descriptions)

        return ParsedDocument(
            name=document_path,
            sections=[
                SectionContent(
                    number=1,
                    text=doc_text,
                    md=doc_text,
                    images=final_images,
                )
            ],
        )
